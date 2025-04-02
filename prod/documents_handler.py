import os
import uuid
import docx
import io
from datetime import datetime
import PyPDF2

class DocumentsHandler:
    """
    Handler for processing uploaded documents.
    Extracts text, chunks it, generates embeddings, and stores in Pinecone.
    """
    def __init__(self, openai_client=None, pinecone_client=None, pinecone_index=None):
        """Initialize the document handler with required clients"""
        self.openai_client = openai_client
        self.pinecone_client = pinecone_client
        self.pinecone_index = pinecone_index
        
    def process_document(self, file_data, filename, chatbot_id, namespace, doc_type="uploaded_doc"):
        """
        Process a document from binary data
        
        Args:
            file_data (bytes): Binary content of the document
            filename (str): Original filename
            chatbot_id (str): ID of the chatbot this document belongs to
            namespace (str): Pinecone namespace for the company
            doc_type (str): Type of document (default: "uploaded_doc")
            
        Returns:
            dict: Processing results including doc_id, vectors_count, chunks, and vectors
        """
        doc_id = str(uuid.uuid4())
        
        # Extract text based on file type
        if filename.lower().endswith('.docx'):
            extracted_text = self.extract_text_from_docx(file_data)
        elif filename.lower().endswith('.doc'):
            # For old .doc files, you might need a different library
            # This is a simplified version that might not work for all .doc files
            extracted_text = self.extract_text_from_docx(file_data)
        elif filename.lower().endswith('.txt'):
            extracted_text = file_data.decode('utf-8')
        elif filename.lower().endswith('.pdf'):
            extracted_text = self.extract_text_from_pdf(file_data)
        else:
            raise ValueError(f"Unsupported file type: {filename}")
        
        # Chunk the text
        chunks = self.chunk_text(extracted_text)
        
        # Generate embeddings
        embeddings = self.get_embeddings(chunks)
        
        # Upload to Pinecone
        vectors_count = self.upload_to_pinecone(
            namespace, 
            chunks, 
            embeddings, 
            doc_id=doc_id
        )
        
        return {
            "doc_id": doc_id,
            "doc_name": filename,
            "chatbot_id": chatbot_id,
            "doc_type": doc_type,
            "content": extracted_text,
            "vectors_count": vectors_count,
            "created_at": datetime.now(),
            "updated_at": datetime.now(),
            # Add these fields for caching
            "chunks": chunks,
            "vectors": embeddings
        }
    
    def extract_text_from_docx(self, file_data):
        """Extract text from a DOCX file"""
        try:
            doc = docx.Document(io.BytesIO(file_data))
            full_text = []
            
            # Extract text from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():  # Skip empty paragraphs
                    full_text.append(para.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        full_text.append(" | ".join(row_text))
            
            return "\n\n".join(full_text)
        except Exception as e:
            print(f"Error extracting text from DOCX: {e}")
            return ""
    
    def extract_text_from_pdf(self, file_data):
        """Extract text from a PDF file"""
        try:
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_data))
            full_text = []
            
            # Get the number of pages in the PDF
            num_pages = len(pdf_reader.pages)
            
            # Extract text from each page
            for page_num in range(num_pages):
                page = pdf_reader.pages[page_num]
                text = page.extract_text()
                
                if text.strip():  # Skip empty pages
                    full_text.append(text)
            
            return "\n\n".join(full_text)
        except Exception as e:
            print(f"Error extracting text from PDF: {e}")
            return ""
    
    def chunk_text(self, text, target_size=700, extended_size=1200):
        """
        Split text into chunks based on semantic boundaries rather than pure character count.
        Preserves structure, keeps URLs with their context, and maintains semantic coherence.
        
        Args:
            text (str): Text to be chunked
            target_size (int): Target size for chunks (default: 700 characters)
            extended_size (int): Extended size limit for special cases (default: 1200 characters)
        
        Returns:
            list: List of text chunks
        """
        # If the entire text is smaller than target_size, return it as a single chunk
        if len(text) <= target_size:
            return [text]
        
        # Split text into initial sections based on double line breaks
        sections = [s.strip() for s in text.split('\n\n') if s.strip()]
        
        chunks = []
        current_chunk = ""
        current_size = 0
        
        for section in sections:
            section_size = len(section)
            contains_url = ('http://' in section or 'https://' in section or 
                            'www.' in section or '.com' in section or '.org' in section)
            
            # Check for list patterns at the start of lines
            list_pattern = any(line.strip().startswith(('- ', '• ', '* ', '· ')) or 
                            (line.strip() and line.strip()[0].isdigit() and line.strip()[1:].startswith('. ')) 
                            for line in section.split('\n'))
            
            # Special handling for sections with URLs or list patterns
            is_special_section = contains_url or list_pattern
            
            # Identify section as a heading (likely to be shorter and be a title/subtitle)
            is_heading = (section_size < 100 and '\n' not in section and 
                        (section.isupper() or section.rstrip().endswith(':') or 
                        (len(section.split()) < 10 and section[-1] not in '.?!')))
            
            # Check if section would put current chunk over extended limit
            if current_size + section_size + 1 > extended_size and current_size > 0:
                # Add current chunk to list and reset
                chunks.append(current_chunk)
                current_chunk = ""
                current_size = 0
            
            # Check if section would fit within target size or deserves extended size
            if (current_size + section_size + 1 <= target_size or 
                (is_special_section and current_size + section_size + 1 <= extended_size) or
                current_size == 0):  # Always include at least one section in a chunk
                
                # Add a separator if not starting a new chunk
                if current_size > 0:
                    current_chunk += "\n\n"
                    current_size += 2
                
                current_chunk += section
                current_size += section_size
            else:
                # Section doesn't fit in current chunk, check if it needs further splitting
                if section_size > extended_size:
                    # Split long section by sentences if possible
                    sentences = []
                    # Simple but relatively effective sentence splitting
                    for sentence in section.replace('!', '.').replace('?', '.').split('.'):
                        if sentence.strip():
                            sentences.append(sentence.strip() + '.')
                    
                    # If we can't split by sentences, fall back to splitting by lines
                    if not sentences:
                        sentences = [line.strip() for line in section.split('\n') if line.strip()]
                    
                    # Add current chunk if we have one
                    if current_chunk:
                        chunks.append(current_chunk)
                        current_chunk = ""
                        current_size = 0
                    
                    # Process sentences to create new chunks
                    temp_chunk = ""
                    temp_size = 0
                    
                    for sentence in sentences:
                        if temp_size + len(sentence) + 1 <= extended_size or temp_size == 0:
                            if temp_size > 0:
                                temp_chunk += " "
                                temp_size += 1
                            temp_chunk += sentence
                            temp_size += len(sentence)
                        else:
                            chunks.append(temp_chunk)
                            temp_chunk = sentence
                            temp_size = len(sentence)
                    
                    if temp_chunk:
                        # If current_chunk is not empty, try to append temp_chunk to it
                        if current_chunk and current_size + temp_size + 2 <= extended_size:
                            current_chunk += "\n\n" + temp_chunk
                            current_size += temp_size + 2
                        else:
                            # Start a new chunk with temp_chunk
                            if current_chunk:
                                chunks.append(current_chunk)
                            current_chunk = temp_chunk
                            current_size = temp_size
                else:
                    # Section is too big for current chunk but not needing further splitting
                    chunks.append(current_chunk)
                    current_chunk = section
                    current_size = section_size
        
        # Add the last chunk if it's not empty
        if current_chunk:
            chunks.append(current_chunk)
        
        return chunks
    
    def get_embeddings(self, text_chunks):
        """Get embeddings for text chunks using OpenAI"""
        embeddings = []
        for chunk in text_chunks:
            response = self.openai_client.embeddings.create(
                input=chunk,
                model="text-embedding-ada-002"
            )
            embeddings.append(response.data[0].embedding)
        return embeddings
    
    def upload_to_pinecone(self, namespace, text_chunks, embeddings, doc_id=None):
        """Upload vectors to Pinecone with document metadata"""
        try:
            index = self.pinecone_client.Index(self.pinecone_index)
            
            vectors = [
                (f"{namespace}-{doc_id}-{i}", 
                 embedding, 
                 {
                     "text": chunk,
                     "doc_id": doc_id,
                     "doc_type": "uploaded",
                     "chunk_index": i
                 })
                for i, (chunk, embedding) in enumerate(zip(text_chunks, embeddings))
            ]
            
            # Upsert vectors in batches of 100
            batch_size = 100
            for i in range(0, len(vectors), batch_size):
                batch = vectors[i:i+batch_size]
                index.upsert(vectors=batch, namespace=namespace)
            
            return len(vectors)
        except Exception as e:
            print(f"Error in Pinecone upload: {e}")
            return 0
    
    def delete_document_vectors(self, namespace, doc_id):
        """Delete all vectors for a specific document from Pinecone"""
        try:
            index = self.pinecone_client.Index(self.pinecone_index)
            
            # Find all vectors with this doc_id
            # Note: This is a simple implementation - in a production system,
            # you might want to store vector IDs for faster deletion
            filter_obj = {"doc_id": {"$eq": doc_id}}
            
            # Delete vectors
            delete_response = index.delete(
                namespace=namespace,
                filter=filter_obj
            )
            
            return True
        except Exception as e:
            print(f"Error deleting document vectors: {e}")
            return False
            
    def create_initial_document_from_website(self, chatbot_id, namespace, processed_content):
        """
        Create the initial document entry from website content
        
        Args:
            chatbot_id (str): ID of the chatbot
            namespace (str): Pinecone namespace
            processed_content (str): The processed website content
            
        Returns:
            dict: Document information
        """
        doc_id = str(uuid.uuid4())
        
        # For initial documents, vectors are already uploaded during website processing
        # So we just need to create the document record
        return {
            "doc_id": doc_id,
            "doc_name": "Scraped Content",
            "chatbot_id": chatbot_id,
            "doc_type": "scraped_content",
            "content": processed_content,
            # We don't know the vector count, would need to query Pinecone
            "vectors_count": 0,
            "created_at": datetime.now(),
            "updated_at": datetime.now()
        }